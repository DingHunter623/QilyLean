#!/usr/bin/env python3
"""Rebuild the QilyLean confidentiality statement signature/date/QR block for clean A4 layout."""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "build" / "nda" / "qilylean-mutual-nda-v1.docx"
QR_PATH = ROOT / "build" / "nda" / "qilylean-trust-qr.png"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_border(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge, values in kwargs.items():
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        for key, value in values.items():
            element.set(qn(f"w:{key}"), str(value))


def set_cell_margins(cell, top=90, start=140, bottom=90, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_cm):
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    width = tc_pr.find(qn("w:tcW"))
    if width is None:
        width = OxmlElement("w:tcW")
        tc_pr.append(width)
    width.set(qn("w:w"), str(int(Cm(width_cm).emu / 635)))
    width.set(qn("w:type"), "dxa")


def style_run(run, size=9.6, bold=False, color="405E5C"):
    run.font.name = "Noto Sans CJK SC"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Noto Sans CJK SC")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_line(paragraph, label, blank, size=9.6):
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.15
    style_run(paragraph.add_run(label), size=size)
    style_run(paragraph.add_run(blank), size=size)


def main():
    if not DOCX_PATH.exists() or not QR_PATH.exists():
        raise SystemExit("Generated DOCX or QR image is missing")

    document = Document(DOCX_PATH)
    if not document.tables:
        raise SystemExit("Signature table is missing")
    old_table = document.tables[-1]

    table = document.add_table(rows=2, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = (6.9, 6.8, 2.9)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    qr_cell = table.cell(0, 2).merge(table.cell(1, 2))

    for text, column in (("委托信息", 0), ("QilyLean｜启力精益", 1)):
        cell = table.cell(0, column)
        set_cell_shading(cell, "EAF5F3")
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(0)
        style_run(paragraph.add_run(text), size=10.5, bold=True, color="0F4B5A")

    left = table.cell(1, 0)
    set_cell_shading(left, "F8FBFA")
    add_line(left.paragraphs[0], "委托单位（填写）：", "____________________")
    add_line(left.add_paragraph(), "项目名称（填写）：", "____________________")

    middle = table.cell(1, 1)
    set_cell_shading(middle, "F8FBFA")
    paragraph = middle.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    style_run(paragraph.add_run("项目责任人：丁启利"))
    add_line(middle.add_paragraph(), "项目责任人签名：", "______________")
    add_line(middle.add_paragraph(), "签署日期：", "______年____月____日")

    set_cell_shading(qr_cell, "F8FBFA")
    paragraph = qr_cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(1)
    style_run(paragraph.add_run("官网核验"), size=9.4, bold=True, color="0F4B5A")
    paragraph = qr_cell.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run().add_picture(str(QR_PATH), width=Cm(2.05), height=Cm(2.05))
    paragraph = qr_cell.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    style_run(paragraph.add_run("qilylean.com"), size=7.7, color="526B69")

    for row_index, row in enumerate(table.rows):
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        for cell in row.cells:
            set_cell_border(
                cell,
                left={"val": "single", "sz": "5", "color": "CADDD9"},
                right={"val": "single", "sz": "5", "color": "CADDD9"},
                top={"val": "single", "sz": "12" if row_index == 0 else "5", "color": "0F7D86" if row_index == 0 else "CADDD9"},
                bottom={"val": "single", "sz": "5", "color": "CADDD9"},
            )

    old_element = old_table._tbl
    old_element.addprevious(table._tbl)
    old_element.getparent().remove(old_element)
    document.save(DOCX_PATH)
    print(f"Optimized signature block in {DOCX_PATH}")


if __name__ == "__main__":
    main()
