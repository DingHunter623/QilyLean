#!/usr/bin/env python3
"""Build the one-page QilyLean confidentiality statement from the approved canonical source."""
from pathlib import Path

import qrcode
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
BUILD = ROOT / "build" / "nda"
BUILD.mkdir(parents=True, exist_ok=True)
DOCX_PATH = BUILD / "qilylean-mutual-nda-v1.docx"
LOGO_PATH = BUILD / "qilylean-document-logo.png"
QR_PATH = BUILD / "qilylean-trust-qr.png"

FONT_BOLD = "/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc"
FONT_REGULAR = "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"


def set_cell_border_none(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        borders.append(element)
    tc_pr.append(borders)


def create_logo():
    image = Image.new("RGBA", (1100, 180), "white")
    draw = ImageDraw.Draw(image)
    font_big = ImageFont.truetype(FONT_BOLD, 82)
    font_small = ImageFont.truetype(FONT_REGULAR, 24)
    font_cn = ImageFont.truetype(FONT_BOLD, 54)

    draw.rounded_rectangle((15, 20, 125, 135), radius=20, fill="#0f4b5a")
    draw.ellipse((43, 42, 72, 71), fill="#ffffff")
    draw.rectangle((48, 70, 67, 110), fill="#ffffff")
    draw.arc((50, 60, 100, 115), start=280, end=100, fill="#ffffff", width=12)
    draw.ellipse((78, 32, 96, 50), fill="#d84a3b")
    draw.text((150, 31), "QilyLean", font=font_big, fill="#0f4b5a")
    draw.line((560, 28, 560, 136), fill="#caa15f", width=5)
    draw.text((585, 51), "启力精益", font=font_cn, fill="#0f4b5a")
    draw.text(
        (150, 136),
        "LEAN MANUFACTURING · ENGINEERING IMPROVEMENT · DIGITAL FACTORY",
        font=font_small,
        fill="#526b69",
    )
    image.save(LOGO_PATH)


def create_qr():
    code = qrcode.QRCode(
        version=3,
        error_correction=qrcode.constants.ERROR_CORRECT_M,
        box_size=8,
        border=2,
    )
    code.add_data("https://qilylean.com/trust/")
    code.make(fit=True)
    code.make_image(fill_color="#0f4b5a", back_color="white").convert("RGB").save(QR_PATH)


def add_font(run, size=None, bold=None, color=None, underline=None):
    run.font.name = "Noto Sans CJK SC"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if underline is not None:
        run.underline = underline
    return run


def build_document():
    create_logo()
    create_qr()

    document = Document()
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(15)
    section.left_margin = Mm(22)
    section.right_margin = Mm(22)
    section.header_distance = Mm(7)
    section.footer_distance = Mm(8)

    normal = document.styles["Normal"]
    normal.font.name = "Noto Sans CJK SC"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
    normal.font.size = Pt(11.2)

    style_specs = (
        ("QL Title", 22, True, "182420", 0, 10, 1.0),
        ("QL Body", 11.2, False, "182420", 0, 4, 1.55),
        ("QL Clause", 10.8, False, "182420", 0, 3, 1.45),
    )
    for name, size, bold, color, before, after, spacing in style_specs:
        style = document.styles[name] if name in document.styles else document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Noto Sans CJK SC"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = spacing

    header = section.header
    table = header.add_table(rows=1, cols=2, width=Mm(166))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Mm(102)
    table.columns[1].width = Mm(64)
    for cell in table.rows[0].cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_border_none(cell)

    paragraph = table.cell(0, 0).paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.add_run().add_picture(str(LOGO_PATH), width=Mm(77))

    paragraph = table.cell(0, 1).paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_font(paragraph.add_run("启精益之智，聚企业之力！"), 10.5, True, "0F4B5A")

    line = header.add_paragraph()
    line.paragraph_format.space_before = Pt(0)
    line.paragraph_format.space_after = Pt(0)
    p_pr = line._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "0F4B5A")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    title = document.add_paragraph(style="QL Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(10)
    title.paragraph_format.space_after = Pt(15)
    title_run = title.add_run("保密声明")
    title_run.font.name = "Noto Serif CJK SC"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Serif CJK SC")
    title_run.font.size = Pt(22)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(24, 36, 32)

    intro = document.add_paragraph(style="QL Body")
    intro.paragraph_format.first_line_indent = Mm(7.4)
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    segments = (
        ("QilyLean｜启力精益（项目责任人：丁启利）受", False),
        ("________________________", True),
        ("委托，开展", False),
        ("________________________", True),
        ("项目。为确保委托单位的商业秘密、技术资料、经营信息及相关人员权益不受侵害，本人／项目组特此声明。", False),
    )
    for text, underline in segments:
        add_font(intro.add_run(text), 11.2, underline=underline)

    heading = document.add_paragraph(style="QL Body")
    heading.paragraph_format.space_before = Pt(2)
    heading.paragraph_format.space_after = Pt(6)
    add_font(heading.add_run("声明内容："), 11.2, True)

    clauses = (
        "本人知悉并同意，对在委托单位调研期间取得的数据资料、产品信息（含未上市产品）、图纸、工艺、设备参数、成本及经营信息严格保密；",
        "本人知悉并同意，对调研期间发现的企业问题点以及拍摄的图片、视频严格保密；如委托单位设有禁止拍照区域，应提前告知QilyLean项目组，未经许可不进行拍摄；",
        "本人知悉并同意，对在委托单位收集的原始书面或电子资料，仅供QilyLean项目内部分析及项目交付使用，使用完毕后按委托单位要求归还、删除或销毁；",
        "本人知悉并同意，对委托单位及其客户的各类商业信息严格保密，不参与委托单位内部利益交换，不与相关人员建立与项目无关的利益关系；",
        "本人知悉并同意，对委托单位内部相关人员的信息（包括但不限于姓名、岗位、联系方式）严格保密；",
        "本人知悉并同意，未经委托单位许可，不携带与项目无关的人员进入委托单位参观、调研或接触相关资料；",
        "本人知悉并同意，QilyLean形成的调研诊断报告、改善方案、成果分享PPT及其他项目资料，未经委托单位书面授权，不向任何第三方提供、转发或公开；经授权用于案例展示的内容，须完成必要的脱敏与去标识化处理。",
    )
    for number, clause in enumerate(clauses, 1):
        paragraph = document.add_paragraph(style="QL Clause")
        paragraph.paragraph_format.left_indent = Mm(7)
        paragraph.paragraph_format.first_line_indent = Mm(-7)
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_font(paragraph.add_run(f"{number}、"), 10.8, True)
        add_font(paragraph.add_run(clause), 10.8)

    signature = document.add_table(rows=2, cols=2)
    signature.alignment = WD_TABLE_ALIGNMENT.CENTER
    signature.autofit = False
    signature.columns[0].width = Mm(96)
    signature.columns[1].width = Mm(70)
    for row in signature.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border_none(cell)

    paragraph = signature.cell(0, 0).paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(0)
    add_font(paragraph.add_run("委托单位（填写）：________________________"), 10.2, color="526B69")

    paragraph = signature.cell(1, 0).paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    add_font(paragraph.add_run("项目名称（填写）：________________________"), 10.2, color="526B69")

    nested = signature.cell(0, 1).add_table(rows=1, cols=2)
    nested.autofit = False
    nested.columns[0].width = Mm(49)
    nested.columns[1].width = Mm(21)
    for cell in nested.rows[0].cells:
        set_cell_border_none(cell)

    paragraph = nested.cell(0, 0).paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    add_font(paragraph.add_run("QilyLean｜启力精益"), 11, True, "0F4B5A")
    add_font(paragraph.add_run("\n项目责任人签名：____________"), 10.5)
    add_font(paragraph.add_run("\n日期：______年____月____日"), 10.5)

    paragraph = nested.cell(0, 1).paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(QR_PATH), width=Mm(17))
    add_font(paragraph.add_run("\n官网核验"), 7)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    add_font(
        paragraph.add_run("QilyLean｜启力精益  ·  https://qilylean.com  ·  电话：134 5001 4003  ·  邮箱：admin@qilylean.com"),
        7.8,
        color="526B69",
    )
    paragraph = footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    add_font(
        paragraph.add_run("文件编号：QL-LEGAL-001  ｜  版本：V1.0  ｜  本范本须结合实际委托主体、项目名称及签署日期填写后使用"),
        7.3,
        color="8D6A32",
    )

    document.core_properties.title = "QilyLean项目保密声明"
    document.core_properties.subject = "制造改善、精益生产、工业工程与数智化工厂项目保密声明范本"
    document.core_properties.author = "QilyLean｜启力精益"
    document.core_properties.keywords = "QilyLean,保密声明,精益改善,工业工程,数智化工厂"
    document.core_properties.comments = "参照单页《保密声明》范本整理为QilyLean版本。"
    document.save(DOCX_PATH)


if __name__ == "__main__":
    build_document()
    print(f"Built {DOCX_PATH}")
